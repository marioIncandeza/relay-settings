
import tkinter as tk
from tkinter import filedialog
import xlwings as xw
from docx import Document
from docx.shared import Pt
from docx.shared import RGBColor
from docx.enum.table import WD_TABLE_ALIGNMENT
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_ALIGN_VERTICAL
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import Inches
from natsort import natsorted
import re

burgundy = [112, 17, 13]
three_col_widths = [1.25, 2.88, 2.88]
std_three_col_names = ['Element', 'Logic', 'Description']
std_tbl_pg_align = WD_TABLE_ALIGNMENT.CENTER
std_tbl_style = 'Grid Table 6 Colorful Accent 3'


def set_cell_background(cell, color):
    """
    Set background color of a table cell.
    :param cell: The cell object.
    :param color: A hex color string, e.g., 'FF0000' for red.
    """
    tc = cell._tc
    tcPr = tc.get_or_add_tcPr()
    shd = OxmlElement('w:shd')
    shd.set(qn('w:fill'), color)
    tcPr.append(shd)


def set_cell_style(cell, style):
    if style == 'white_buc':
        cell.paragraphs[0].runs[0].font.color.rgb = RGBColor(255, 255, 255)  # white
        set_cell_background(cell, '5C2A2B')  # buccaneer
        cell.paragraphs[0].runs[0].bold = True


def set_table(docx_table, filtered_settings, header_names, font_color, col_widths):
    set_header_row(docx_table, header_names, font_color, col_widths)
    for i, row in enumerate(filtered_settings):
        row_cells = docx_table.rows[i + 1].cells
        for j, value in enumerate(row):
            row_cells[j].text = str(value)
            # Apply formatting from Excel (e.g., font size)
            row_cells[j].width = Inches(col_widths[j])
            row_cells[j].paragraphs[0].runs[0].bold = False
            row_cells[j].paragraphs[0].runs[0].font.size = Pt(12.0)
            row_cells[j].paragraphs[0].runs[0].font.color.rgb = RGBColor(font_color[0], font_color[1], font_color[2])
            row_cells[j].vertical_alignment = WD_ALIGN_VERTICAL.CENTER
            row_cells[j].paragraphs[0].alignment = WD_ALIGN_PARAGRAPH.LEFT


def set_table_no_headers(docx_table, filtered_settings, font_color, col_widths):
    for i, row in enumerate(filtered_settings):
        row_cells = docx_table.rows[i].cells
        for j, value in enumerate(row):
            if value:
                row_cells[j].text = str(value)
            else:
                row_cells[j].text = ''
            # Apply formatting from Excel (e.g., font size)
            row_cells[j].width = Inches(col_widths[j])
            row_cells[j].paragraphs[0].runs[0].bold = False
            row_cells[j].paragraphs[0].runs[0].font.size = Pt(12.0)
            row_cells[j].paragraphs[0].runs[0].font.color.rgb = RGBColor(font_color[0], font_color[1], font_color[2])
            row_cells[j].vertical_alignment = WD_ALIGN_VERTICAL.CENTER
            row_cells[j].paragraphs[0].alignment = WD_ALIGN_PARAGRAPH.LEFT


def set_table_title(doc, table_title):
    # First row
    table = doc.add_table(rows=1, cols=1)
    table.autofit = False
    first_row = table.rows[0].cells
    first_row[0].text = table_title
    first_row[0].paragraphs[0].runs[0].font.color.rgb = RGBColor(255, 255, 255)  # white
    set_cell_background(first_row[0], '5c2a2b')
    for paragraph in first_row[0].paragraphs:
        paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER
    first_row[0].vertical_alignment = WD_ALIGN_VERTICAL.CENTER
    first_row[0].paragraphs[0].runs[0].bold = True
    first_row[0].paragraphs[0].paragraph_format.line_spacing = 1
    first_row[0].paragraphs[0].paragraph_format.space_after = 0


def set_vfreq_table_title(doc):
    table = doc.add_table(rows=2, cols=6)
    table.autofit = False
    table.style = 'burgundy_white_border'
    # Description
    table.cell(0, 0).merge(table.cell(1, 0))
    table.cell(0, 0).text = 'Description'
    table.cell(0, 0).paragraphs[0].runs[0].font.color.rgb = RGBColor(255, 255, 255) # white
    table.cell(0, 0).width = Inches(2)
    # Pickup
    table.cell(0, 1).merge(table.cell(0, 2))
    table.cell(0, 1).text = 'Pickup'
    table.cell(0, 1).paragraphs[0].runs[0].font.color.rgb = RGBColor(255, 255, 255)  # white
    table.cell(0, 1).paragraphs[0].alignment = WD_ALIGN_PARAGRAPH.CENTER
    table.cell(0, 1).width = Inches(2)
    # Element
    table.cell(1, 1).text = 'Element'
    table.cell(1, 1).paragraphs[0].runs[0].font.color.rgb = RGBColor(255, 255, 255)  # white
    table.cell(1, 1).paragraphs[0].alignment = WD_ALIGN_PARAGRAPH.CENTER
    table.cell(1, 1).width = Inches(1)
    # Value
    table.cell(1, 2).text = 'Value'
    table.cell(1, 2).paragraphs[0].runs[0].font.color.rgb = RGBColor(255, 255, 255)  # white
    table.cell(1, 2).paragraphs[0].alignment = WD_ALIGN_PARAGRAPH.CENTER
    table.cell(1, 2).width = Inches(1)
    # Logic
    table.cell(0, 3).merge(table.cell(1, 3))
    table.cell(0, 3).text = 'Logic'
    table.cell(0, 3).paragraphs[0].runs[0].font.color.rgb = RGBColor(255, 255, 255)  # white
    table.cell(0, 3).paragraphs[0].alignment = WD_ALIGN_PARAGRAPH.CENTER
    table.cell(0, 3).width = Inches(1)
    # Pickup Delay
    table.cell(0, 4).merge(table.cell(0, 5))
    table.cell(0, 4).text = 'Pickup Delay'
    table.cell(0, 4).paragraphs[0].runs[0].font.color.rgb = RGBColor(255, 255, 255)  # white
    table.cell(0, 4).paragraphs[0].alignment = WD_ALIGN_PARAGRAPH.CENTER
    table.cell(0, 4).width = Inches(2)
    # Seconds
    table.cell(1, 4).text = 'Seconds'
    table.cell(1, 4).paragraphs[0].runs[0].font.color.rgb = RGBColor(255, 255, 255)  # white
    table.cell(1, 4).paragraphs[0].alignment = WD_ALIGN_PARAGRAPH.CENTER
    table.cell(1, 4).width = Inches(1)
    # Cycles
    table.cell(1, 5).text = 'Cycles'
    table.cell(1, 5).paragraphs[0].runs[0].font.color.rgb = RGBColor(255, 255, 255)  # white
    table.cell(1, 5).paragraphs[0].alignment = WD_ALIGN_PARAGRAPH.CENTER
    table.cell(1, 5).width = Inches(1)


def filter_settings(settings_table, category, filter_column='Category', all_columns=False):
    category_idx = settings_table[0].index(filter_column)
    filtered_settings = []
    for i, row in enumerate(settings_table, start=1):
        if row[category_idx] == category:
            if all_columns:
                filtered_settings.append(row)
            else:
                filtered_settings.append(row[:3])
    return filtered_settings


def reorg_sel_vars(filtered_settings):
    sel_vars = []
    for row in filtered_settings:
        if 'SV' in row[0]:
            if 'PU' not in row[0] and 'DO' not in row[0]:
                temp = [row[0], row[1], None, None, row[2]]
                sel_vars.append(temp)

    for row in sel_vars:
        for settings in filtered_settings:
            if row[0] in settings[0] and 'PU' in settings[0]:
                row[2] = settings[1]
            if row[0] in settings[0] and 'DO' in settings[0]:
                row[3] = settings[1]
        if row[2] is None:
            row[2] = 0.00
        if row[3] is None:
            row[3] = 0.00

    return sel_vars


def reorg_latch(filtered_settings):
    latch_bits = []
    latches = []
    sets = []
    resets = []
    descrips = {}
    for row in filtered_settings:
        if ('LT' + row[0][3:]) not in latches:
            latch_name = 'LT' + row[0][3:]
            latches.append(latch_name)
            descrips[latch_name] = row[2]
        if 'SET' in row[0]:
            sets.append(row[1])
        if 'RST' in row[0]:
            resets.append(row[1])
    latches = natsorted(latches)
    sets = natsorted(sets)
    resets = natsorted(resets)
    for i, latch in enumerate(latches):
        temp = [latch, sets[i], resets[i], descrips[latch]]
        latch_bits.append(temp)
    return latch_bits


def reorg_display(filtered_settings):
    display_points = []
    elements = []
    logic = {}
    set_messages = []
    clear_messages = []
    descrips = {}
    for row in filtered_settings:
        if '_' not in row[0]:
            elements.append(row[0])
            logic[row[0]] = row[1]
            descrips[row[0]] = row[2]
        if '_1' in row[0]:
            set_messages.append(row[1])
        if '_0' in row[0]:
            clear_messages.append(row[1])
    elements = natsorted(elements)
    set_messages = natsorted(set_messages)
    clear_messages = natsorted(clear_messages)
    for i, element in enumerate(elements):
        temp = [element, logic[element], set_messages[i], clear_messages[i], descrips[element]]
        display_points.append(temp)
    return display_points


def reorg_vfreq_summary(filtered_settings):
    summary = []
    underfrequency_block = []
    descrips = {}
    elements = []
    values = {}
    svs = {}
    cycs = {}
    freq_elements = []
    volt_elements = []
    for row in filtered_settings:
        if row[0] == '27B81P':
            underfrequency_block = [row[2], row[0], row[1], '', '', '']
        if '81D' in row[0] and 'P' in row[0]:
            freq_elements.append(row[0])
            values[row[0]] = row[1]
            descrips[row[0]] = row[2]
            for inside_row in filtered_settings:
                if row[0][:4] == inside_row[0][:4] and 'D' in inside_row[0]:
                    cycs[row[0]] = inside_row[1]
            svs[row[0]] = ''
        if '27P' in row[0]:
            volt_elements.append(row[0])
            values[row[0]] = row[1]
            descrips[row[0]] = row[2]
            for inside_row in filtered_settings:
                if 'SV' in inside_row[0] and re.match('27.' + row[0][3], str(inside_row[1])):
                    svs[row[0]] = inside_row[0]
                    for inside_row_2 in filtered_settings:
                        if inside_row[0] in inside_row_2[0] and 'PU' in inside_row_2[0]:
                            cycs[row[0]] = inside_row_2[1]
        if '59P' in row[0]:
            volt_elements.append(row[0])
            values[row[0]] = row[1]
            descrips[row[0]] = row[2]
            for inside_row in filtered_settings:
                if 'SV' in inside_row[0] and re.match('59.' + row[0][3], str(inside_row[1])):
                    svs[row[0]] = inside_row[0]
                    for inside_row_2 in filtered_settings:
                        if inside_row[0] in inside_row_2[0] and 'PU' in inside_row_2[0]:
                            cycs[row[0]] = inside_row_2[1]
    freq_elements = natsorted(freq_elements)
    volt_elements = natsorted(volt_elements)
    summary.append(underfrequency_block)
    elements.extend(freq_elements)
    elements.extend(volt_elements)
    for i, element in enumerate(elements):
        temp = [descrips[element], element, values[element], svs[element], str(round(float(cycs[element])/60,3)), cycs[element]]
        summary.append(temp)
    return summary


def set_header_row(docx_table, header_names, header_font_color, header_col_widths, bold=True):
    for i, name in enumerate(header_names):
        active_row = docx_table.rows[0].cells
        active_row[i].text = name
        active_row[i].paragraphs[0].runs[0].font.color.rgb = RGBColor(header_font_color[0], header_font_color[1], header_font_color[2])  # burgundy
        active_row[i].paragraphs[0].runs[0].bold = True
        active_row[i].width = Inches(header_col_widths[i])


def create_settings_group_table(docx, title, column_count, style, page_align, autofit, filtered_settings, col_names,
                                font_color, col_widths):
    set_table_title(docx, title)
    table = docx.add_table(rows=len(filtered_settings)+1, cols=column_count) # +1 for column descriptions
    table.style = style
    table.alignment = page_align
    table.autofit = autofit

    set_table(table, filtered_settings, col_names, font_color, col_widths)


def create_vfreq_table(docx, column_count, style, page_align, autofit, filtered_settings,
                       font_color, col_widths):
    set_vfreq_table_title(docx)
    table = docx.add_table(rows=len(filtered_settings), cols=column_count)
    table.style = style
    table.alignment = page_align
    table.autofit = autofit

    set_table_no_headers(table, filtered_settings, font_color, col_widths)


def create_io_table(docx, column_count, style, page_align, autofit, io, font_color, col_widths):
    table = docx.add_table(rows=len(io), cols=column_count)
    table.style = style
    table.alignment = page_align
    table.autofit = autofit

    set_table_no_headers(table, io, font_color, col_widths)

    for i in range(column_count):
        set_cell_style(table.cell(0, i), 'white_buc')


def create_oc_table(docx, column_count, style, page_align, autofit, data, font_color, col_widths):
    table = docx.add_table(rows=len(data), cols=column_count)
    table.style = style
    table.alignment = page_align
    table.autofit = autofit

    set_table_no_headers(table, data, font_color, col_widths)

    for i, row in enumerate(table.rows):
        row_cells = row.cells
        row_cells[0].paragraphs[0].runs[0].bold = True
        if i == 0:
            for cell in row_cells:
                if cell.paragraphs[0].runs:
                    cell.paragraphs[0].runs[0].bold = True
                cell.paragraphs[0].alignment = WD_ALIGN_PARAGRAPH.CENTER
            temp = row_cells[0].text
            row_cells[0].merge(row_cells[1])
            row_cells[0].text = temp
            row_cells[0].paragraphs[0].runs[0].font.color.rgb = RGBColor(112, 17, 13)
        if i == 3:
            for k, cell in enumerate(row_cells):
                if k != 0:
                    set_cell_style(cell, 'white_buc')
        for cell in row_cells:
            cell.paragraphs[0].alignment = WD_ALIGN_PARAGRAPH.CENTER


def gen_351S():
    try:
        # Open the Excel file and select the sheet
        root = tk.Tk()
        root.withdraw()

        app = xw.App(visible=False)
        xl_path = filedialog.askopenfilename(title="Select a file")
        wb = app.books.open(xl_path)
        sheet = wb.sheets['FeederLogic']

        # Read the settings
        io_351S = sheet.tables['io_351S'].range.value  # Access named table range directly
        main_settings_351S = sheet.tables['settings_351S'].range.value  # Access named table range directly
        ph_oc_1 = sheet.tables['PH_OC_1'].range.value  # Access named table range directly
        ph_oc_2 = sheet.tables['PH_OC_2'].range.value  # Access named table range directly

        # Create a Word document
        doc = Document(
            r"C:\Users\laerps\OneDrive - Westwood Active Directory\Desktop\351S\Relay Settings Report - Stripped.docx")

        doc.add_heading('11F Input/Output Summary', level=2)
        create_io_table(doc, 2, std_tbl_style, WD_TABLE_ALIGNMENT.LEFT, True, io_351S, burgundy, [2, 2])
        doc.add_paragraph()

        doc.add_heading('Feeder PH OC Settings', level=2)
        create_oc_table(doc, 4, std_tbl_style, WD_TABLE_ALIGNMENT.LEFT, True, ph_oc_1, burgundy, [1, 1, 1, 1])
        doc.add_paragraph()
        create_oc_table(doc, 4, std_tbl_style, WD_TABLE_ALIGNMENT.LEFT, True, ph_oc_2, burgundy, [1, 1, 1, 1])
        doc.add_paragraph()

        doc.add_heading('Voltage and Frequency', level=3)
        filtered_settings = filter_settings(main_settings_351S, 'V/Freq Protection', filter_column='Function',
                                            all_columns=True)
        volt_freq_settings = reorg_vfreq_summary(filtered_settings)
        create_vfreq_table(doc, 6, std_tbl_style, std_tbl_pg_align, True, volt_freq_settings, burgundy, [2, 1, 1, 1, 1, 1])

        # Add a heading to section
        doc.add_heading('11F Logic', level=2)
        doc.add_paragraph(text='The table below outlines the major settings used in the feeder relay.')

        filtered_settings = filter_settings(main_settings_351S, 'Global')
        create_settings_group_table(doc, 'Global', 3, std_tbl_style, std_tbl_pg_align,
                                    False, filtered_settings, std_three_col_names, burgundy, three_col_widths)

        filtered_settings = filter_settings(main_settings_351S, 'Trip Logic')
        create_settings_group_table(doc, 'Trip Logic', 3, std_tbl_style, std_tbl_pg_align,
                                    False, filtered_settings, std_three_col_names, burgundy, three_col_widths)

        filtered_settings = filter_settings(main_settings_351S, 'SELogic Variables')
        sel_vars = reorg_sel_vars(filtered_settings)
        create_settings_group_table(doc, 'SELogic Variables', 5, std_tbl_style, std_tbl_pg_align, False,
                                    sel_vars, ['Element', 'Logic', 'PU (cyc)', 'DO (cyc)', 'Description'],
                                    burgundy, [0.75, 2.5, 0.75, 0.75, 2.25])

        filtered_settings = filter_settings(main_settings_351S, 'Latch Bits')
        latch_bits = reorg_latch(filtered_settings)
        create_settings_group_table(doc, 'Latch Bits', 4, std_tbl_style, std_tbl_pg_align, False,
                                    latch_bits, ['Element', 'Set', 'Reset', 'Description'],
                                    burgundy, [0.75, 2, 2, 2.25])

        filtered_settings = filter_settings(main_settings_351S, 'Display Points')
        display_points = reorg_display(filtered_settings)
        create_settings_group_table(doc, 'Display Points', 5, std_tbl_style, std_tbl_pg_align, False,
                                    display_points, ['Element', 'Logic', 'Set Message', 'Clear Message', 'Description'],
                                    burgundy, [0.75, 1.33, 1.33, 1.33, 2.25])
        # Save the document
        save_path = filedialog.asksaveasfilename(title="Save file as",
                                                 defaultextension=".docx",
                                                 filetypes=[("Word documents", "*.docx"), ("All files", "*.*")]
                                                 )
        doc.save(save_path)

    # Close workbook and quit app
    finally:
        wb.close()
        app.quit()


gen_351S()